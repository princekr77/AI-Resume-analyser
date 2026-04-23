"""
Resume parsing utilities for extracting text from PDF and DOCX files
"""

import io
import re
from typing import Optional
import PyPDF2
import docx
import fitz  # PyMuPDF

class ResumeParser:
    """Handle extraction of text from various resume formats"""
    
    @staticmethod
    def extract_text_from_pdf(file) -> Optional[str]:
        """
        Extract text from PDF file using PyMuPDF (best accuracy)
        
        Args:
            file: Uploaded PDF file object
            
        Returns:
            Extracted text as string, or None if extraction fails
        """
        try:
            # Use PyMuPDF for better text extraction
            pdf_document = fitz.open(stream=file.read(), filetype="pdf")
            text = ""
            
            for page_num in range(pdf_document.page_count):
                page = pdf_document[page_num]
                text += page.get_text()
            
            pdf_document.close()
            file.seek(0)  # Reset file pointer
            
            # Clean up the extracted text
            text = ResumeParser._clean_text(text)
            return text
            
        except Exception as e:
            print(f"Error extracting PDF with PyMuPDF: {e}")
            
            # Fallback to PyPDF2
            try:
                file.seek(0)
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text()
                
                text = ResumeParser._clean_text(text)
                file.seek(0)
                return text
            except Exception as e2:
                print(f"Error extracting PDF with PyPDF2: {e2}")
                return None
    
    @staticmethod
    def extract_text_from_docx(file) -> Optional[str]:
        """
        Extract text from DOCX file
        
        Args:
            file: Uploaded DOCX file object
            
        Returns:
            Extracted text as string, or None if extraction fails
        """
        try:
            doc = docx.Document(file)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += "\n" + cell.text
            
            text = ResumeParser._clean_text(text)
            file.seek(0)
            return text
            
        except Exception as e:
            print(f"Error extracting DOCX: {e}")
            return None
    
    @staticmethod
    def _clean_text(text: str) -> str:
        """
        Clean and normalize extracted text
        
        Args:
            text: Raw extracted text
            
        Returns:
            Cleaned text
        """
        if not text:
            return ""
        
        # Remove extra whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters but keep important punctuation
        text = re.sub(r'[^\w\s\.,;:!?@#$%&*()\-+=<>]', '', text)
        
        # Normalize line breaks
        text = text.replace('\r', '\n')
        text = re.sub(r'\n\s*\n', '\n\n', text)
        
        return text.strip()
    
    @staticmethod
    def extract_text(file) -> Optional[str]:
        """
        Extract text from resume file based on file type
        
        Args:
            file: Uploaded file object
            
        Returns:
            Extracted text as string, or None if extraction fails
        """
        filename = file.name.lower()
        
        if filename.endswith('.pdf'):
            return ResumeParser.extract_text_from_pdf(file)
        elif filename.endswith('.docx'):
            return ResumeParser.extract_text_from_docx(file)
        else:
            raise ValueError("Unsupported file format. Please upload PDF or DOCX files.")